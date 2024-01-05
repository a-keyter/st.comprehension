import io
import requests
import random
import os

import streamlit as st

from pytube import YouTube

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from langchain.document_loaders import YoutubeLoader

from langchain import LLMChain
from langchain.chat_models import ChatOpenAI
from langchain.prompts.chat import (
    ChatPromptTemplate,
    SystemMessagePromptTemplate,
    HumanMessagePromptTemplate,
)

############ STREAMLIT APP ##############

with st.sidebar:
    st.title("AI Listening Comprehension")
    st.header("Turn YouTube Videos into Listening Comprehension Exercises")
    st.write("Enter a link to a youtube video, to produce a listening comprehension worksheet")
    st.divider()
    user_openai = st.text_input("Please enter your OpenAI API Key:",)

st.title("AI Listening Comprehension")

user_link = st.text_input("Enter a link to a YouTube Video")
target_lang = st.selectbox("Select the original language of the video", ["English", "Spanish", "Spanish - LATAM", "French", "German", "Italian", "Portuguese",])

lang_codes = {
    "English": "en",
    "Spanish": "es",
    "Spanish - LATAM": "es-419",
    "French": "fr",
    "German": "de",
    "Italian": "it",
    "Portuguese": "pt",
}

generate = st.button("Create Worksheet")

if generate and user_openai == "":
  st.error("You must enter an OpenAI API Key")
elif generate and user_openai != "":
  #Set API keys
  os.environ['OPENAI_API_KEY'] = user_openai
  chat_model = "gpt-4-1106-preview"
  #Set Temp
  temp = 0.7

  def get_youtube_details(url):
      yt = YouTube(url)
      title = yt.title
      thumbnail_url = yt.thumbnail_url
      return title, thumbnail_url

  def key_vocab(vid_transcript, target_language):
      chat = ChatOpenAI(temperature=temp, model=chat_model) # Here temperature is set to temp to provide a balanced response

      template = f"""
      You are a {target_language} language teacher asked to prepare reading and listening comprehension tasks from a youtube video. 
      
      You will be given a transcript from a youtube video and you should return a list of 12 key vocabulary terms from the video in {target_language}.
    
      """

      system_message_prompt = SystemMessagePromptTemplate.from_template(template)

      human_template = f""""

      Here is the transcript of a video in {target_language}:

      {vid_transcript}

      Give me 12 key vocabulary terms (Including their definite article) from the video transcript. Include the definite article for each noun if the language has gendered nouns. Seperate each term with #NEWTERM#.

      """

      human_message_prompt = HumanMessagePromptTemplate.from_template(human_template)
      chat_prompt = ChatPromptTemplate.from_messages([system_message_prompt, human_message_prompt])

      chain = LLMChain(llm=chat, prompt=chat_prompt)
      result = chain.run({"vid_transcript": vid_transcript, "target_language": target_language})
      return result #

  def def_vocab(vocabulary, target_language):
      chat = ChatOpenAI(temperature=temp, model=chat_model) # Here temperature is set to temp to provide a balanced response

      template = f"""
      You are a {target_language} language teacher asked to prepare reading and listening comprehension tasks. 
      
      You will be given a list of vocabulary terms. For each of the vocabulary terms you will return a definition of each term using simple language in {target_language}. Seperate each definition with: #NEWDEF#
    
      """

      system_message_prompt = SystemMessagePromptTemplate.from_template(template)

      human_template = f""""

      Here is the vocabulary in {target_language}:

      {vocabulary}

      Give me a definition of each term using simple language in {target_language}. Do not use the term in the response. Start each definition with the {target_language} translation of "This...". Seperate each definition with: #NEWDEF#

      """

      human_message_prompt = HumanMessagePromptTemplate.from_template(human_template)
      chat_prompt = ChatPromptTemplate.from_messages([system_message_prompt, human_message_prompt])

      chain = LLMChain(llm=chat, prompt=chat_prompt)
      result = chain.run({"vocabulary": vocabulary, "target_language": target_language})
      return result #

  def comprehension_q(vid_transcript, target_language):
      chat = ChatOpenAI(temperature=temp, model=chat_model) # Here temperature is set to temp to provide a balanced response

      template = f"""
      You are a {target_language} language teacher asked to prepare reading and listening comprehension tasks from a youtube script. You will be given a transcript from a youtube video and you should return a series of comprehension questions in the {target_language} to check if a students has understood the video. 

      """

      system_message_prompt = SystemMessagePromptTemplate.from_template(template)

      human_template = f"""

      Here is the transcript of a video in {target_language}:

      {vid_transcript}

      Give me 6 comprehension questions about the video ranging from simple to difficult. Start each question with #QUESTION#. After each question, give the answer to the question. Start each answer with #ANSWER#. Do not repeat any questions.

      """

      human_message_prompt = HumanMessagePromptTemplate.from_template(human_template)
      chat_prompt = ChatPromptTemplate.from_messages([system_message_prompt, human_message_prompt])

      chain = LLMChain(llm=chat, prompt=chat_prompt)
      result = chain.run({"vid_transcript": vid_transcript, "target_language": target_language})
      return result #

  def download_image(url):
      response = requests.get(url)
      if response.status_code == 200:
          return io.BytesIO(response.content)
      return None

  loader = YoutubeLoader.from_youtube_url(
    user_link,
    add_video_info=False,
    language=lang_codes[target_lang],
  )

  try:
    transcript = loader.load()
  except:
      st.error("Transcript for this video is not available, please try a different video")

  title, thumbnail_url = get_youtube_details(user_link)

  st.subheader(title)
  st.image(thumbnail_url, caption='YouTube Video Thumbnail')

  with st.spinner("Now creating your worksheet"):
    # Get Key Vocab
    vocab = key_vocab(vid_transcript=str(transcript[0].page_content), target_language=target_lang)

    vocab_list = vocab.split("#NEWTERM#")

    def strip_enumeration(word_list):
      return [word.split('. ', 1)[1] if '. ' in word else word for word in word_list]

    vocab_list = strip_enumeration(vocab_list)

    with st.expander("Key Vocab"):
      st.write(vocab_list)

    # Generate Vocab Definitions
    vocab_defs = def_vocab(vocab_list, target_lang)
    vocab_defs.strip()
    vocab_definition_list = vocab_defs.split("#NEWDEF#")
    vocab_definition_list = [entry.strip() for entry in vocab_definition_list]

    with st.expander("Vocab definitions:"):
      st.write(vocab_definition_list)

    # Generate Comprehension Questions
    questions = comprehension_q(vid_transcript = str(transcript[0].page_content), target_language= target_lang)
    
    qa_list = questions.split('#QUESTION#')

    qa_pairs = []

    for qa in qa_list:
      if qa.strip() != "":
        question, answer = qa.split("#ANSWER#")
        qa_pairs.append({"question": question.strip(), "answer": answer.strip()})
    
    with st.expander("Questions about the video"):
      for qa in qa_pairs:
        st.write(qa["question"])
      for qa in qa_pairs:
        st.write(qa["answer"])

    worksheet = Document()

    # Define narrow margins
    horizontal_margin = Inches(1)
    vertical_margin = Inches(0.5)

    # Set margins for each section in the document
    for section in worksheet.sections:
        section.top_margin = vertical_margin
        section.bottom_margin = vertical_margin
        section.left_margin = horizontal_margin
        section.right_margin = horizontal_margin

    paragraph_format = worksheet.styles['Normal'].paragraph_format
    paragraph_format.space_after = None

    worksheet.add_heading(title, level=1)
    head = worksheet.paragraphs[-1]
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_format = head.paragraph_format
    heading_format.space_before = Pt(0)  # Adjust space before heading
    heading_format.space_after = Pt(5)  # Adjust space after heading

    image_stream = download_image(thumbnail_url)
    if image_stream:
      worksheet.add_picture(image_stream, width=Inches(4))
      thumb = worksheet.paragraphs[-1]
      thumb.alignment = WD_ALIGN_PARAGRAPH.CENTER
      thumb_format = thumb.paragraph_format
      thumb_format.space_after = Pt(10)  # Adjust space after heading

    worksheet.add_heading("Vocabulary", level=2)
    table = worksheet.add_table(rows=4, cols=3)

    # Populate the table with vocab
    for i, vocab in enumerate(vocab_list):
      row = i // 3
      col = i % 3
      paragraph = table.cell(row, col).paragraphs[0]
      paragraph.add_run(f"{i+1}. {vocab}")

    worksheet.add_heading("Match the Vocabulary to the Definitions:", level=2)
    shuffled_definitions = vocab_definition_list[:]
    random.shuffle(shuffled_definitions)
    for definition in shuffled_definitions:
      worksheet.add_paragraph(f"__________ - {definition}")

    worksheet.add_heading("Comprehension Questions", level=2)

    for qa in qa_pairs:
      worksheet.add_paragraph(qa["question"])
      answer_line1 = worksheet.add_paragraph('_' * 100)
      answer_line2 = worksheet.add_paragraph('_' * 100)
      answer_line_format = answer_line1.paragraph_format
      answer_line_format.space_before = Pt(5)  # Adjust space before line
      answer_line_format.space_after = Pt(10)  # Adjust space after line
      answer_line_format = answer_line2.paragraph_format
      answer_line_format.space_before = Pt(5)  # Adjust space before line
      answer_line_format.space_after = Pt(10)  # Adjust space after line

    worksheet.add_page_break()
    worksheet.add_heading("Worksheet Answers", level=2)

    for i in range(10):
      worksheet.add_paragraph(f"{vocab_list[i]} - {vocab_definition_list[i]}")
    
    for qa in qa_pairs:
      question = qa["question"]
      answer = qa["answer"]
      question_answer = question + " " + answer
      worksheet.add_paragraph(question_answer)
    
    bio = io.BytesIO()
    worksheet.save(bio)

  st.success("Successfully Generated Worksheet")

  st.download_button(
      label="Download Worksheet",
      data=bio.getvalue(),
      file_name=f"Listening Comprehension - {title}.docx",
      mime="docx",
  )


